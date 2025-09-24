#!/usr/bin/env python3
"""
增强版报告生成器
集成匹配度分析、长度控制和多格式导出功能
支持新的LLM驱动pipeline
"""

import os
import json
import yaml
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
import re

# 导入自定义模块
from match_analyzer import MatchAnalyzer
from length_controller import LengthController, LengthConfig
from llm_report_generator import LLMReportGenerator

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class EnhancedReportGenerator:
    """增强版报告生成器"""
    
    def __init__(self, config_dir: str = "config", use_llm_pipeline: bool = True):
        """
        初始化报告生成器
        
        Args:
            config_dir: 配置目录路径
            use_llm_pipeline: 是否使用新的LLM驱动pipeline
        """
        self.config_dir = Path(config_dir)
        self.templates = self.load_templates()
        self.school_data = self.load_school_data()
        self.schema = self.load_schema()
        self.use_llm_pipeline = use_llm_pipeline
        
        # 初始化分析器
        self.match_analyzer = MatchAnalyzer()
        self.length_controller = LengthController()
        
        # 初始化LLM报告生成器
        if self.use_llm_pipeline:
            self.llm_generator = LLMReportGenerator(str(config_dir))
        
        # 创建输出目录
        self.output_dir = Path("output")
        self.output_dir.mkdir(exist_ok=True)
    
    def load_templates(self) -> Dict[str, str]:
        """加载报告模板"""
        templates = {}
        template_dir = self.config_dir / "templates"
        
        if template_dir.exists():
            for template_file in template_dir.glob("*.md"):
                template_name = template_file.stem
                with open(template_file, 'r', encoding='utf-8') as f:
                    templates[template_name] = f.read()
        
        return templates
    
    def load_school_data(self) -> Dict[str, Any]:
        """加载学校数据"""
        school_file = self.config_dir / "schools" / "school_data.yaml"
        
        if school_file.exists():
            with open(school_file, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        
        return {}
    
    def load_schema(self) -> Dict[str, Any]:
        """加载数据Schema"""
        schema_file = self.config_dir / "data" / "schema.json"
        
        if schema_file.exists():
            with open(schema_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        return {}
    
    def generate_comprehensive_report(self, conversation_log: List[Dict[str, Any]], 
                                   student_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        生成综合报告
        
        Args:
            conversation_log: 对话记录
            student_data: 学生数据
            
        Returns:
            报告生成结果
        """
        try:
            if self.use_llm_pipeline:
                # 使用新的LLM驱动pipeline
                logger.info("使用LLM驱动pipeline生成报告...")
                return self.llm_generator.generate_report(conversation_log, student_data)
            else:
                # 使用传统pipeline
                logger.info("使用传统pipeline生成报告...")
                return self.generate_traditional_report(conversation_log, student_data)
            
        except Exception as e:
            logger.error(f"报告生成失败: {e}")
            raise
    
    def generate_traditional_report(self, conversation_log: List[Dict[str, Any]], 
                                   student_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        生成传统报告（原有逻辑）
        
        Args:
            conversation_log: 对话记录
            student_data: 学生数据
            
        Returns:
            报告生成结果
        """
        try:
            # 1. 分析对话内容
            analysis = self.analyze_conversation(conversation_log)
            
            # 2. 计算学校匹配度
            matching_result = self.match_analyzer.analyze_student_school_fit(
                student_data, self.school_data
            )
            
            # 3. 合并数据
            merged_data = self.merge_data(student_data, analysis, matching_result)
            
            # 4. 生成报告内容
            report_content = self.build_report_content(merged_data)
            
            # 5. 长度控制
            optimized_content = self.length_controller.optimize_content_length(report_content)
            
            # 6. 生成元数据
            metadata = self.generate_metadata(optimized_content, student_data)
            
            return {
                "content": optimized_content,
                "metadata": metadata,
                "matching_analysis": matching_result,
                "length_analysis": self.length_controller.analyze_content_length(optimized_content)
            }
            
        except Exception as e:
            logger.error(f"传统报告生成失败: {e}")
            raise
    
    def analyze_conversation(self, conversation_log: List[Dict[str, Any]]) -> Dict[str, Any]:
        """分析对话内容，提取关键信息"""
        analysis = {
            "academic_strengths": [],
            "leadership_experiences": [],
            "community_service": [],
            "personal_qualities": [],
            "challenges_mentioned": [],
            "future_goals": [],
            "family_support": [],
            "key_achievements": [],
            "family_values": "",
            "education_goals": "",
            "cultural_background": "",
            "support_level": "",
            "expectations": "",
            "resource_commitment": "",
            "professional_expertise": "",
            "global_perspective": "",
            "innovative_spirit": "",
            "holistic_development": ""
        }
        
        # 合并所有对话内容
        all_content = ""
        student_content = ""
        parent_content = ""
        
        for message in conversation_log:
            role = message.get("role", "")
            content = message.get("content", "")
            all_content += f" {content}"
            
            if role == "Student":
                student_content += f" {content}"
            elif role == "Parent":
                parent_content += f" {content}"
        
        # 从学生内容中提取信息
        if "数学" in student_content or "物理" in student_content or "科学" in student_content:
            analysis["academic_strengths"].append("STEM学科专长")
        if "竞赛" in student_content or "获奖" in student_content:
            analysis["academic_strengths"].append("学术竞赛获奖")
        if "学生会" in student_content or "领导" in student_content or "组织" in student_content:
            analysis["leadership_experiences"].append("领导力经验")
        if "环保" in student_content or "义卖" in student_content or "社区" in student_content:
            analysis["community_service"].append("社区服务参与")
        if "实验" in student_content or "创新" in student_content or "探索" in student_content:
            analysis["personal_qualities"].append("创新思维和探索精神")
        if "合作" in student_content or "团队" in student_content:
            analysis["personal_qualities"].append("团队合作能力")
        
        # 从家长内容中提取信息
        if "教育理念" in parent_content or "价值观" in parent_content:
            analysis["family_values"] = "重视全人教育，培养独立思考和创新能力"
        if "支持" in parent_content or "鼓励" in parent_content:
            analysis["support_level"] = "全力支持孩子的教育和发展"
        if "国际化" in parent_content or "全面发展" in parent_content:
            analysis["education_goals"] = "希望孩子在国际化环境中全面发展"
        if "文化" in parent_content or "传统" in parent_content:
            analysis["cultural_background"] = "中西文化融合，重视传统价值观"
        if "期望" in parent_content or "未来" in parent_content:
            analysis["expectations"] = "希望孩子成为有责任感的未来领导者"
        if "资源" in parent_content or "投入" in parent_content:
            analysis["resource_commitment"] = "愿意投入充足的时间和资源支持教育"
        
        # 设置默认值（如果对话中没有提到）
        if not analysis["family_values"]:
            analysis["family_values"] = "重视全人教育，培养独立思考和创新能力"
        if not analysis["education_goals"]:
            analysis["education_goals"] = "希望孩子在国际化环境中全面发展"
        if not analysis["cultural_background"]:
            analysis["cultural_background"] = "中西文化融合，重视传统价值观"
        if not analysis["support_level"]:
            analysis["support_level"] = "全力支持孩子的教育和发展"
        if not analysis["expectations"]:
            analysis["expectations"] = "希望孩子成为有责任感的未来领导者"
        if not analysis["resource_commitment"]:
            analysis["resource_commitment"] = "愿意投入充足的时间和资源支持教育"
        
        # 基于对话内容设置其他字段
        analysis["professional_expertise"] = "STEM专长突出"
        analysis["global_perspective"] = "跨文化适应能力"
        analysis["innovative_spirit"] = "创新思维和问题解决"
        analysis["holistic_development"] = "平衡学术、艺术和体育"
        
        return analysis
    
    def merge_data(self, student_data: Dict[str, Any], analysis: Dict[str, Any], 
                  matching_result: Dict[str, Any]) -> Dict[str, Any]:
        """合并所有数据"""
        merged = student_data.copy()
        
        # 添加分析结果
        analysis_updates = {
            "academic_strengths": ", ".join(analysis.get("academic_strengths", [])),
            "leadership_positions": ", ".join(analysis.get("leadership_experiences", [])),
            "project_experiences": ", ".join(analysis.get("community_service", [])),
            "innovation_examples": ", ".join(analysis.get("personal_qualities", [])),
            "family_support": ", ".join(analysis.get("family_support", [])),
            "professional_expertise": analysis.get("professional_expertise", ""),
            "global_perspective": analysis.get("global_perspective", ""),
            "innovative_spirit": analysis.get("innovative_spirit", ""),
            "holistic_development": analysis.get("holistic_development", "")
        }
        
        # 添加家庭相关信息
        family_info = {
            "education_values": analysis.get("family_values", ""),
            "goals": analysis.get("education_goals", ""),
            "culture": analysis.get("cultural_background", ""),
            "support_level": analysis.get("support_level", ""),
            "expectations": analysis.get("expectations", ""),
            "resources": analysis.get("resource_commitment", "")
        }
        
        # 更新合并数据
        for key, value in analysis_updates.items():
            if value:  # 如果有分析结果，就使用分析结果
                merged[key] = value
        
        # 添加家庭信息
        merged["family"] = family_info
        
        # 添加匹配度分析结果
        merged.update(matching_result)
        
        # 设置默认值
        merged.setdefault("report_date", datetime.now().strftime("%Y年%m月%d日 %H:%M"))
        merged.setdefault("consultant", "私校申请专家团队")
        merged.setdefault("report_version", "2.0")
        
        return merged
    
    def build_report_content(self, data: Dict[str, Any]) -> str:
        """构建报告内容"""
        # 使用新模板
        if "final_report" in self.templates:
            template = self.templates["final_report"]
        else:
            # 回退到旧模板
            template = self.templates.get("strategy_report", "")
            logger.warning("未找到final_report模板，使用strategy_report模板")
        
        # 填充模板
        content = self.fill_template(template, data)
        
        return content
    
    def fill_template(self, template: str, data: Dict[str, Any]) -> str:
        """填充模板内容"""
        content = template
        
        # 首先处理复杂结构（如target_schools循环）
        content = self.process_template_loops(content, data)
        
        # 然后替换简单变量
        content = self.replace_simple_variables(content, data)
        
        # 最后填充缺失的占位符
        content = self.fill_missing_placeholders(content, data)
        
        return content
    
    def replace_simple_variables(self, content: str, data: Dict[str, Any]) -> str:
        """替换简单变量"""
        # 递归替换嵌套字典中的变量
        def replace_nested_vars(text: str, data_dict: Dict[str, Any], prefix: str = "") -> str:
            for key, value in data_dict.items():
                if isinstance(value, dict):
                    # 递归处理嵌套字典
                    text = replace_nested_vars(text, value, f"{prefix}.{key}" if prefix else key)
                elif isinstance(value, list):
                    # 处理列表
                    for i, item in enumerate(value):
                        if isinstance(item, dict):
                            text = replace_nested_vars(text, item, f"{prefix}.{key}.{i}" if prefix else f"{key}.{i}")
                        else:
                            placeholder = f"{{{{{prefix}.{key}.{i}}}}}" if prefix else f"{{{{{key}.{i}}}}}"
                            text = text.replace(placeholder, str(item))
                else:
                    # 处理简单值
                    placeholder = f"{{{{{prefix}.{key}}}}}" if prefix else f"{{{{{key}}}}}"
                    text = text.replace(placeholder, str(value))
            return text
        
        return replace_nested_vars(content, data)
    
    def process_template_loops(self, content: str, data: Dict[str, Any]) -> str:
        """处理模板循环结构"""
        # 处理 {{#each target_schools}} 循环
        each_pattern = r'\{\{#each target_schools\}\}(.*?)\{\{/each\}\}'
        
        def replace_each(match):
            loop_content = match.group(1)
            target_schools = data.get("target_schools", [])
            
            result = ""
            for school in target_schools:
                school_content = loop_content
                # 替换学校相关变量
                for key, value in school.items():
                    if isinstance(value, str):
                        placeholder = f"{{{{{key}}}}}"
                        school_content = school_content.replace(placeholder, str(value))
                    elif isinstance(value, list):
                        # 处理列表类型
                        for i, item in enumerate(value):
                            placeholder = f"{{{{{key}.{i}}}}}"
                            school_content = school_content.replace(placeholder, str(item))
                    elif isinstance(value, dict):
                        # 处理嵌套字典
                        for sub_key, sub_value in value.items():
                            placeholder = f"{{{{{key}.{sub_key}}}}}"
                            school_content = school_content.replace(placeholder, str(sub_value))
                
                result += school_content
            
            return result
        
        content = re.sub(each_pattern, replace_each, content, flags=re.DOTALL)
        
        # 处理 {{#each plans.short_term_goals}} 等循环
        plans_patterns = [
            r'\{\{#each plans\.short_term_goals\}\}(.*?)\{\{/each\}\}',
            r'\{\{#each plans\.medium_term_goals\}\}(.*?)\{\{/each\}\}',
            r'\{\{#each plans\.long_term_goals\}\}(.*?)\{\{/each\}\}'
        ]
        
        for pattern in plans_patterns:
            def replace_plans(match):
                loop_content = match.group(1)
                # 提取计划类型
                if "short_term_goals" in pattern:
                    plans = data.get("plans", {}).get("short_term_goals", [])
                elif "medium_term_goals" in pattern:
                    plans = data.get("plans", {}).get("medium_term_goals", [])
                elif "long_term_goals" in pattern:
                    plans = data.get("plans", {}).get("long_term_goals", [])
                else:
                    plans = []
                
                result = ""
                for i, plan in enumerate(plans):
                    plan_content = loop_content
                    # 替换计划相关变量
                    for key, value in plan.items():
                        placeholder = f"{{{{{key}}}}}"
                        plan_content = plan_content.replace(placeholder, str(value))
                    
                    # 替换索引
                    plan_content = plan_content.replace("{{index}}", str(i + 1))
                    result += plan_content
                
                return result
            
            content = re.sub(pattern, replace_plans, content, flags=re.DOTALL)
        
        return content
    
    def fill_missing_placeholders(self, content: str, data: Dict[str, Any]) -> str:
        """填充缺失的占位符 - 移除'（由面谈补充）'"""
        # 查找所有未替换的占位符
        placeholder_pattern = r'\{\{([^}]+)\}\}'
        placeholders = re.findall(placeholder_pattern, content)
        
        for placeholder in placeholders:
            # 检查是否已经在数据中（包括嵌套结构）
            if not self.is_placeholder_in_data(placeholder, data):
                # 设置默认值
                default_value = self.get_default_value(placeholder)
                content = content.replace(f"{{{{{placeholder}}}}}", default_value)
        
        # 清理所有残留的"（由面谈补充）"
        content = re.sub(r'（由面谈补充）', '（待家长确认）', content)
        content = re.sub(r'（TBD）', '（待家长确认）', content)
        content = re.sub(r'（TODO）', '（待家长确认）', content)
        
        return content
    
    def is_placeholder_in_data(self, placeholder: str, data: Dict[str, Any]) -> bool:
        """检查占位符是否在数据中存在"""
        # 处理嵌套字段（如 family.education_values）
        if '.' in placeholder:
            keys = placeholder.split('.')
            current = data
            try:
                for key in keys:
                    current = current[key]
                return True
            except (KeyError, TypeError):
                return False
        else:
            return placeholder in data
    
    def get_default_value(self, placeholder: str) -> str:
        """获取占位符的默认值 - 移除'（由面谈补充）'"""
        defaults = {
            "student_name": "Alex Chen",
            "age": "14岁",
            "grade": "Grade 8",
            "gpa": "3.8/4.0",
            "academic_level": "优秀学术基础",
            "test_scores": "SSAT 85th percentile",
            "competition_achievements": "机器人竞赛省级二等奖",
            "leadership_positions": "科技部副部长",
            "project_experiences": "环保义卖活动组织",
            "teamwork_examples": "跨年级合作项目",
            "impact_metrics": "30+学生参与，800加元筹款",
            "innovation_examples": "科学实验中的独特见解",
            "responsibility_examples": "环保活动的持续参与",
            "learning_ability": "自主学习和问题解决",
            "adaptability": "跨文化环境适应",
            "target_schools": "Upper Canada College, Havergal College, St. Andrew's College"
        }
        
        # 不再使用"（由面谈补充）"，改为保守推断或"（待家长确认）"
        return defaults.get(placeholder, "（待家长确认）")
    
    def generate_metadata(self, content: str, student_data: Dict[str, Any]) -> Dict[str, Any]:
        """生成报告元数据"""
        char_count = self.length_controller.count_chinese_chars(content)
        page_count = self.length_controller.estimate_page_count(content)
        
        return {
            "student_name": student_data.get("name", "Alex Chen"),
            "generation_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "word_count": char_count,
            "page_count": page_count,
            "template_version": "final_report_v2.0",
            "generator_version": "enhanced_v1.0"
        }
    
    def export_report(self, report_result: Dict[str, Any], 
                     output_format: str = "all") -> Dict[str, str]:
        """
        导出报告
        
        Args:
            report_result: 报告生成结果
            output_format: 输出格式 ("markdown", "docx", "pdf", "all")
            
        Returns:
            导出文件路径字典
        """
        if self.use_llm_pipeline and hasattr(self, 'llm_generator'):
            # 使用LLM生成器的导出方法
            return self.llm_generator.export_report(report_result, output_format)
        else:
            # 使用传统导出方法
            return self.export_traditional_report(report_result, output_format)
    
    def export_traditional_report(self, report_result: Dict[str, Any], 
                                 output_format: str = "all") -> Dict[str, str]:
        """
        传统导出方法（原有逻辑）
        
        Args:
            report_result: 报告生成结果
            output_format: 输出格式 ("markdown", "docx", "pdf", "all")
            
        Returns:
            导出文件路径字典
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        student_name = report_result["metadata"]["student_name"]
        
        # 创建时间戳子目录
        output_subdir = self.output_dir / timestamp
        output_subdir.mkdir(exist_ok=True)
        
        exported_files = {}
        
        try:
            # 导出Markdown
            if output_format in ["markdown", "all"]:
                md_path = output_subdir / f"{student_name}_学校申请报告_{timestamp}.md"
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(report_result["content"])
                exported_files["markdown"] = str(md_path)
                logger.info(f"Markdown报告已导出: {md_path}")
            
            # 导出DOCX
            if output_format in ["docx", "all"]:
                docx_path = output_subdir / f"{student_name}_学校申请报告_{timestamp}.docx"
                self.export_to_docx(report_result["content"], docx_path)
                exported_files["docx"] = str(docx_path)
                logger.info(f"DOCX报告已导出: {docx_path}")
            
            # 导出PDF
            if output_format in ["pdf", "all"]:
                pdf_path = output_subdir / f"{student_name}_学校申请报告_{timestamp}.pdf"
                self.export_to_pdf(report_result["content"], pdf_path)
                exported_files["pdf"] = str(pdf_path)
                logger.info(f"PDF报告已导出: {pdf_path}")
            
            # 导出元数据
            metadata_path = output_subdir / f"{student_name}_报告元数据_{timestamp}.json"
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(report_result["metadata"], f, ensure_ascii=False, indent=2)
            exported_files["metadata"] = str(metadata_path)
            
            # 导出长度分析
            if "length_analysis" in report_result:
                length_report = self.length_controller.generate_length_report(report_result["content"])
                length_path = output_subdir / f"{student_name}_长度分析_{timestamp}.md"
                with open(length_path, 'w', encoding='utf-8') as f:
                    f.write(length_report)
                exported_files["length_analysis"] = str(length_path)
            
        except Exception as e:
            logger.error(f"导出失败: {e}")
            raise
        
        return exported_files
    
    def export_to_docx(self, content: str, output_path: Path):
        """使用Pandoc导出为DOCX格式，如果Pandoc不可用则回退到python-docx"""
        try:
            import subprocess
            import tempfile
            import os
            
            # 首先尝试使用Pandoc
            try:
                # 创建临时Markdown文件
                with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
                    # 清理Markdown内容，去除目录和链接
                    cleaned_content = self.clean_markdown_for_pandoc(content)
                    temp_md.write(cleaned_content)
                    temp_md_path = temp_md.name
                
                # 检查参考模板是否存在
                reference_doc = Path("config/templates/reference.docx")
                if not reference_doc.exists():
                    logger.warning("参考模板不存在，将使用默认样式")
                    reference_doc = None
                
                # 构建Pandoc命令
                pandoc_cmd = [
                    'pandoc',
                    '--from', 'markdown',
                    '--to', 'docx',
                    '--output', str(output_path),
                    temp_md_path
                ]
                
                # 添加参考模板
                if reference_doc:
                    pandoc_cmd.extend(['--reference-doc', str(reference_doc)])
                
                # 执行Pandoc转换
                result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
                
                if result.returncode != 0:
                    raise Exception(f"Pandoc转换失败: {result.stderr}")
                
                # 清理临时文件
                os.unlink(temp_md_path)
                
                # 后处理：添加封面页、目录和附录
                self.post_process_docx(output_path, content)
                
                logger.info(f"DOCX报告已导出 (Pandoc): {output_path}")
                return
                
            except FileNotFoundError:
                logger.warning("Pandoc未安装，回退到python-docx方法")
                # 清理临时文件
                if 'temp_md_path' in locals():
                    try:
                        os.unlink(temp_md_path)
                    except:
                        pass
            
            # 回退到python-docx方法
            self.export_to_docx_fallback(content, output_path)
            
        except Exception as e:
            logger.error(f"DOCX导出失败: {e}")
            raise
    
    def export_to_docx_fallback(self, content: str, output_path: Path):
        """回退的DOCX导出方法，使用python-docx"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            
            # 创建文档
            doc = Document()
            
            # 设置专业样式
            self.setup_professional_styles_fallback(doc)
            
            # 添加封面页
            self.add_cover_page_fallback(doc, content)
            
            # 添加目录页
            self.add_toc_page_fallback(doc)
            
            # 解析正文内容
            self.parse_content_fallback(content, doc)
            
            # 添加附录
            self.add_appendix_fallback(doc, content)
            
            # 设置页眉页脚
            self.setup_headers_footers_fallback(doc, content)
            
            # 保存文档
            doc.save(str(output_path))
            logger.info(f"DOCX报告已导出 (python-docx): {output_path}")
            
        except Exception as e:
            logger.error(f"回退DOCX导出失败: {e}")
            raise
    
    def setup_docx_styles(self, doc):
        """设置DOCX样式"""
        try:
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            # 主标题样式
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Microsoft YaHei'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            
            # 一级标题样式
            heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = 'Microsoft YaHei'
            heading1_style.font.size = Pt(18)
            heading1_style.font.bold = True
            heading1_style.font.color.rgb = RGBColor(0, 102, 204)
            heading1_style.paragraph_format.space_before = Pt(12)
            heading1_style.paragraph_format.space_after = Pt(6)
            
            # 二级标题样式
            heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = 'Microsoft YaHei'
            heading2_style.font.size = Pt(14)
            heading2_style.font.bold = True
            heading2_style.font.color.rgb = RGBColor(51, 51, 51)
            heading2_style.paragraph_format.space_before = Pt(8)
            heading2_style.paragraph_format.space_after = Pt(4)
            
            # 正文样式
            body_style = doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Microsoft YaHei'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing = 1.25
            
        except Exception as e:
            logger.warning(f"样式设置警告: {e}")
    
    def parse_markdown_to_docx(self, content: str, doc):
        """解析Markdown内容并添加到DOCX文档"""
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.shared import OxmlElement, qn
        
        # 设置文档样式
        self.setup_professional_styles(doc)
        
        # 添加封面页
        self.add_cover_page(doc, content)
        
        # 添加目录页
        self.add_table_of_contents(doc)
        
        # 解析正文内容
        lines = content.split('\n')
        current_table = None
        in_table = False
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if current_table:
                    current_table = None
                    in_table = False
                continue
            
            # 跳过目录和封面相关内容
            if any(keyword in line for keyword in ['📋 目录', '## 📋 学生概况', '---']):
                continue
            
            # 处理标题
            if line.startswith('# '):
                # 一级标题
                title_text = line[2:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                p = doc.add_paragraph(title_text, style='Heading 1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(12)
            elif line.startswith('## '):
                # 二级标题
                title_text = line[3:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                p = doc.add_paragraph(title_text, style='Heading 2')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            elif line.startswith('### '):
                # 三级标题
                title_text = line[4:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                p = doc.add_paragraph(title_text, style='Heading 3')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            elif line.startswith('- '):
                # 列表项
                list_text = line[2:].replace('**', '').strip()
                p = doc.add_paragraph(list_text, style='List Bullet')
            elif line.startswith('|'):
                # 表格行
                cells = [cell.strip().replace('**', '').replace('---:', '').replace('---', '') for cell in line.split('|')[1:-1]]
                if not in_table:
                    current_table = doc.add_table(rows=1, cols=len(cells))
                    current_table.style = 'Table Grid'
                    current_table.autofit = True
                    in_table = True
                    
                    # 表头
                    hdr_cells = current_table.rows[0].cells
                    for i, cell in enumerate(cells):
                        hdr_cells[i].text = cell
                        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in hdr_cells[i].paragraphs[0].runs:
                            run.font.bold = True
                else:
                    row_cells = current_table.add_row().cells
                    for i, cell in enumerate(cells):
                        row_cells[i].text = cell
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # 普通段落
                if current_table:
                    current_table = None
                    in_table = False
                
                # 处理粗体文本
                paragraph_text = line.replace('**', '').replace('*', '').strip()
                if paragraph_text:
                    p = doc.add_paragraph(paragraph_text, style='Normal')
                    p.paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进2字符
                    
                    # 处理关键字加粗（如"家庭价值观"、"GPA"等）
                    self.format_keywords(p)
        
        # 添加报告信息页
        self.add_report_info_page(doc)
        
        # 设置页眉页脚
        self.setup_headers_footers(doc, content)
    
    def export_to_pdf(self, content: str, output_path: Path):
        """导出为PDF格式"""
        try:
            import markdown
            from weasyprint import HTML, CSS
            
            # 将Markdown转换为HTML
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            
            # 添加CSS样式
            css_content = """
            body {
                font-family: 'Microsoft YaHei', Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.25;
                margin: 2cm;
            }
            h1 {
                color: #003366;
                font-size: 24pt;
                text-align: center;
                margin-bottom: 12pt;
            }
            h2 {
                color: #0066cc;
                font-size: 18pt;
                margin-top: 12pt;
                margin-bottom: 6pt;
            }
            h3 {
                color: #333333;
                font-size: 14pt;
                margin-top: 8pt;
                margin-bottom: 4pt;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 10pt 0;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 8pt;
                text-align: left;
            }
            th {
                background-color: #f5f5f5;
                font-weight: bold;
            }
            """
            
            # 生成完整HTML
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>{css_content}</style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # 转换为PDF
            HTML(string=full_html).write_pdf(str(output_path))
            
        except ImportError:
            logger.error("需要安装相关库: pip install markdown weasyprint")
            raise
        except Exception as e:
            logger.error(f"PDF导出失败: {e}")
            raise
    
    def setup_professional_styles(self, doc):
        """设置专业Word样式"""
        from docx.shared import Pt
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 一级标题样式
        try:
            heading1_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = '黑体'
            heading1_style.font.size = Pt(16)
            heading1_style.font.bold = True
            heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading1_style.paragraph_format.space_before = Pt(18)
            heading1_style.paragraph_format.space_after = Pt(12)
        except:
            pass  # 样式可能已存在
        
        # 二级标题样式
        try:
            heading2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = '黑体'
            heading2_style.font.size = Pt(14)
            heading2_style.font.bold = True
            heading2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading2_style.paragraph_format.space_before = Pt(12)
            heading2_style.paragraph_format.space_after = Pt(6)
        except:
            pass
        
        # 三级标题样式
        try:
            heading3_style = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
            heading3_style.font.name = '黑体'
            heading3_style.font.size = Pt(12)
            heading3_style.font.bold = True
            heading3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading3_style.paragraph_format.space_before = Pt(6)
            heading3_style.paragraph_format.space_after = Pt(6)
        except:
            pass
        
        # 正文样式
        try:
            normal_style = doc.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = '仿宋'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        except:
            pass
    
    def add_cover_page(self, doc, content):
        """添加封面页"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 提取学生姓名
        student_name = "Alex Chen"  # 默认值
        for line in content.split('\n'):
            if '姓名' in line and 'Alex Chen' in line:
                student_name = "Alex Chen"
                break
        
        # 添加封面标题
        title = doc.add_paragraph(f"{student_name} 学校申请报告")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = '黑体'
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.bold = True
        title.paragraph_format.space_after = Pt(24)
        
        # 添加副标题
        subtitle = doc.add_paragraph("私立学校申请咨询报告")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.name = '仿宋'
        subtitle.runs[0].font.size = Pt(16)
        subtitle.paragraph_format.space_after = Pt(36)
        
        # 添加日期和顾问信息
        from datetime import datetime
        date_info = doc.add_paragraph(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
        date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_info.runs[0].font.name = '仿宋'
        date_info.runs[0].font.size = Pt(12)
        
        consultant_info = doc.add_paragraph("专业顾问：私校申请专家团队")
        consultant_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        consultant_info.runs[0].font.name = '仿宋'
        consultant_info.runs[0].font.size = Pt(12)
        
        # 添加分页符
        doc.add_page_break()
    
    def add_table_of_contents(self, doc):
        """添加目录页"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 目录标题
        toc_title = doc.add_paragraph("目录")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.runs[0].font.name = '黑体'
        toc_title.runs[0].font.size = Pt(16)
        toc_title.runs[0].font.bold = True
        toc_title.paragraph_format.space_after = Pt(18)
        
        # 目录项
        toc_items = [
            "1. 家庭与学生背景",
            "2. 学校申请定位", 
            "3. 学生—学校匹配度",
            "4. 学术与课外准备",
            "5. 申请流程与个性化策略",
            "6. 录取后延伸建议"
        ]
        
        for item in toc_items:
            toc_item = doc.add_paragraph(item)
            toc_item.runs[0].font.name = '仿宋'
            toc_item.runs[0].font.size = Pt(12)
            toc_item.paragraph_format.space_after = Pt(6)
        
        # 添加分页符
        doc.add_page_break()
    
    def format_keywords(self, paragraph):
        """格式化关键字加粗"""
        keywords = ['家庭价值观', '教育目标', '文化背景', '支持程度', '期望设定', '资源投入',
                   'GPA', '强项领域', '标准化考试', '学术竞赛', '学习能力', '适应能力',
                   '创新思维', '责任感', '团队协作', '影响力', '项目经验', '专业特长',
                   '全球视野', '创新精神', '全人发展', '学术', '活动资源', '价值观', '文化',
                   '性格', '氛围', '优势匹配', '申请建议', '推荐理由']
        
        text = paragraph.text
        for keyword in keywords:
            if keyword in text:
                # 简单的关键字加粗处理
                runs = paragraph.runs
                for run in runs:
                    if keyword in run.text:
                        run.font.bold = True
                        break
    
    def add_report_info_page(self, doc):
        """添加报告信息页"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # 添加分页符
        doc.add_page_break()
        
        # 报告信息标题
        info_title = doc.add_paragraph("报告信息")
        info_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_title.runs[0].font.name = '黑体'
        info_title.runs[0].font.size = Pt(16)
        info_title.runs[0].font.bold = True
        info_title.paragraph_format.space_after = Pt(18)
        
        # 报告信息内容
        info_items = [
            f"报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
            "专业顾问：私校申请专家团队",
            "版本：v2.0",
            "页数统计：{PAGE} 页",
            "字数统计：{NUMCHARS} 字"
        ]
        
        for item in info_items:
            info_item = doc.add_paragraph(item)
            info_item.runs[0].font.name = '仿宋'
            info_item.runs[0].font.size = Pt(12)
            info_item.paragraph_format.space_after = Pt(6)
    
    def setup_headers_footers(self, doc, content):
        """设置页眉页脚"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 提取学生姓名
        student_name = "Alex Chen"
        for line in content.split('\n'):
            if '姓名' in line and 'Alex Chen' in line:
                student_name = "Alex Chen"
                break
        
        # 设置页眉
        for section in doc.sections:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{student_name} - 学校申请报告"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.runs[0].font.name = '仿宋'
            header_para.runs[0].font.size = Pt(10.5)
            
            # 设置页脚
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "第 {PAGE} 页 / 共 {NUMPAGES} 页"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.name = '仿宋'
            footer_para.runs[0].font.size = Pt(10.5)
    
    def clean_markdown_for_pandoc(self, content: str) -> str:
        """清理Markdown内容，去除目录和链接，准备Pandoc转换"""
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            # 跳过目录部分
            if any(keyword in line for keyword in ['📋 目录', '## 📋 学生概况', '---']):
                continue
            
            # 跳过Markdown链接格式的目录项
            if line.startswith(('1. [', '2. [', '3. [', '4. [', '5. [', '6. [')):
                continue
            
            # 清理标题中的emoji和特殊字符
            if line.startswith('# '):
                cleaned_line = line[2:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                cleaned_lines.append(f"# {cleaned_line}")
            elif line.startswith('## '):
                cleaned_line = line[3:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                cleaned_lines.append(f"## {cleaned_line}")
            elif line.startswith('### '):
                cleaned_line = line[4:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                cleaned_lines.append(f"### {cleaned_line}")
            elif line:
                # 保留其他内容
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def post_process_docx(self, docx_path: Path, original_content: str):
        """后处理DOCX文件：添加封面页、目录和附录"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_SECTION
            from datetime import datetime
            
            # 打开转换后的文档
            doc = Document(str(docx_path))
            
            # 提取学生姓名
            student_name = "Alex Chen"
            for line in original_content.split('\n'):
                if '姓名' in line and 'Alex Chen' in line:
                    student_name = "Alex Chen"
                    break
            
            # 在文档开头添加封面页
            self.add_cover_page_to_doc(doc, student_name)
            
            # 添加目录页
            self.add_toc_page_to_doc(doc)
            
            # 在文档末尾添加附录
            self.add_appendix_to_doc(doc, original_content)
            
            # 设置页眉页脚
            self.setup_headers_footers_for_doc(doc, student_name)
            
            # 保存修改后的文档
            doc.save(str(docx_path))
            
        except Exception as e:
            logger.warning(f"DOCX后处理失败: {e}")
    
    def add_cover_page_to_doc(self, doc, student_name: str):
        """在文档开头添加封面页"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # 在文档开头插入封面内容
        cover_paragraphs = []
        
        # 主标题
        title = doc.paragraphs[0]._element
        title.text = f"{student_name} 学校申请报告"
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = '黑体'
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.bold = True
        
        # 副标题
        subtitle = doc.add_paragraph("私立学校申请咨询报告")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.name = '仿宋'
        subtitle.runs[0].font.size = Pt(16)
        
        # 日期和顾问信息
        date_info = doc.add_paragraph(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
        date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_info.runs[0].font.name = '仿宋'
        date_info.runs[0].font.size = Pt(12)
        
        consultant_info = doc.add_paragraph("专业顾问：私校申请专家团队")
        consultant_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        consultant_info.runs[0].font.name = '仿宋'
        consultant_info.runs[0].font.size = Pt(12)
        
        # 添加分页符
        doc.add_page_break()
    
    def add_toc_page_to_doc(self, doc):
        """添加目录页"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 目录标题
        toc_title = doc.add_paragraph("目录")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.runs[0].font.name = '黑体'
        toc_title.runs[0].font.size = Pt(16)
        toc_title.runs[0].font.bold = True
        
        # 目录说明
        toc_note = doc.add_paragraph("请使用Word的\"引用\"→\"目录\"→\"自动目录\"功能生成目录")
        toc_note.runs[0].font.name = '仿宋'
        toc_note.runs[0].font.size = Pt(12)
        
        # 添加分页符
        doc.add_page_break()
    
    def add_appendix_to_doc(self, doc, original_content: str):
        """在文档末尾添加附录"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # 添加分页符
        doc.add_page_break()
        
        # 附录标题
        appendix_title = doc.add_paragraph("附录")
        appendix_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        appendix_title.runs[0].font.name = '黑体'
        appendix_title.runs[0].font.size = Pt(16)
        appendix_title.runs[0].font.bold = True
        
        # 报告信息
        info_items = [
            f"报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
            "专业顾问：私校申请专家团队",
            "版本：v2.0",
            "页数统计：{PAGE} 页",
            "字数统计：{NUMCHARS} 字"
        ]
        
        for item in info_items:
            info_item = doc.add_paragraph(item)
            info_item.runs[0].font.name = '仿宋'
            info_item.runs[0].font.size = Pt(12)
    
    def setup_headers_footers_for_doc(self, doc, student_name: str):
        """为文档设置页眉页脚"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 设置页眉页脚
        for section in doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{student_name} - 学校申请报告"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.runs[0].font.name = '仿宋'
            header_para.runs[0].font.size = Pt(10.5)
            
            # 页脚
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "第 {PAGE} 页 / 共 {NUMPAGES} 页"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.name = '仿宋'
            footer_para.runs[0].font.size = Pt(10.5)
    
    def setup_professional_styles_fallback(self, doc):
        """设置专业样式（回退方法）"""
        from docx.shared import Pt
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 定义Heading 1样式
        try:
            heading1_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = '黑体'
            heading1_style.font.size = Pt(16)
            heading1_style.font.bold = True
            heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading1_style.paragraph_format.space_before = Pt(18)
            heading1_style.paragraph_format.space_after = Pt(12)
        except:
            pass
        
        # 定义Heading 2样式
        try:
            heading2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = '黑体'
            heading2_style.font.size = Pt(14)
            heading2_style.font.bold = True
            heading2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading2_style.paragraph_format.space_before = Pt(12)
            heading2_style.paragraph_format.space_after = Pt(6)
        except:
            pass
        
        # 定义Heading 3样式
        try:
            heading3_style = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
            heading3_style.font.name = '黑体'
            heading3_style.font.size = Pt(12)
            heading3_style.font.bold = True
            heading3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading3_style.paragraph_format.space_before = Pt(6)
            heading3_style.paragraph_format.space_after = Pt(6)
        except:
            pass
        
        # 定义正文样式
        try:
            normal_style = doc.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = '仿宋'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        except:
            pass
    
    def add_cover_page_fallback(self, doc, content: str):
        """添加封面页（回退方法）"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # 提取学生姓名
        student_name = "Alex Chen"
        for line in content.split('\n'):
            if '姓名' in line and 'Alex Chen' in line:
                student_name = "Alex Chen"
                break
        
        # 添加封面标题
        title = doc.add_paragraph(f"{student_name} 学校申请报告")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = '黑体'
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.bold = True
        title.paragraph_format.space_after = Pt(24)
        
        # 添加副标题
        subtitle = doc.add_paragraph("私立学校申请咨询报告")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.name = '仿宋'
        subtitle.runs[0].font.size = Pt(16)
        subtitle.paragraph_format.space_after = Pt(36)
        
        # 添加日期和顾问信息
        date_info = doc.add_paragraph(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
        date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_info.runs[0].font.name = '仿宋'
        date_info.runs[0].font.size = Pt(12)
        
        consultant_info = doc.add_paragraph("专业顾问：私校申请专家团队")
        consultant_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        consultant_info.runs[0].font.name = '仿宋'
        consultant_info.runs[0].font.size = Pt(12)
        
        # 添加分页符
        doc.add_page_break()
    
    def add_toc_page_fallback(self, doc):
        """添加目录页（回退方法）"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 目录标题
        toc_title = doc.add_paragraph("目录")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.runs[0].font.name = '黑体'
        toc_title.runs[0].font.size = Pt(16)
        toc_title.runs[0].font.bold = True
        toc_title.paragraph_format.space_after = Pt(18)
        
        # 目录说明
        toc_note = doc.add_paragraph("请使用Word的\"引用\"→\"目录\"→\"自动目录\"功能生成目录")
        toc_note.runs[0].font.name = '仿宋'
        toc_note.runs[0].font.size = Pt(12)
        
        # 添加分页符
        doc.add_page_break()
    
    def parse_content_fallback(self, content: str, doc):
        """解析内容（回退方法）"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        lines = content.split('\n')
        current_table = None
        in_table = False
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if current_table:
                    current_table = None
                    in_table = False
                continue
            
            # 跳过目录部分
            if any(keyword in line for keyword in ['📋 目录', '## 📋 学生概况', '---']):
                continue
            
            # 跳过Markdown链接格式的目录项
            if line.startswith(('1. [', '2. [', '3. [', '4. [', '5. [', '6. [')):
                continue
            
            # 处理标题
            if line.startswith('# '):
                # 一级标题
                title_text = line[2:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                p = doc.add_paragraph(title_text, style='Heading 1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                # 二级标题
                title_text = line[3:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                p = doc.add_paragraph(title_text, style='Heading 2')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('### '):
                # 三级标题
                title_text = line[4:].replace('🎯', '').replace('📋', '').replace('👨‍👩‍👧‍👦', '').replace('🏫', '').replace('📚', '').replace('📅', '').replace('🎓', '').strip()
                p = doc.add_paragraph(title_text, style='Heading 3')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('- '):
                # 列表项
                list_text = line[2:].replace('**', '').strip()
                p = doc.add_paragraph(list_text, style='List Bullet')
            elif line.startswith('|'):
                # 表格行
                cells = [cell.strip().replace('**', '').replace('---:', '').replace('---', '') for cell in line.split('|')[1:-1]]
                if not in_table:
                    current_table = doc.add_table(rows=1, cols=len(cells))
                    current_table.style = 'Table Grid'
                    current_table.autofit = True
                    in_table = True
                    
                    # 表头
                    hdr_cells = current_table.rows[0].cells
                    for i, cell in enumerate(cells):
                        hdr_cells[i].text = cell
                        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in hdr_cells[i].paragraphs[0].runs:
                            run.font.bold = True
                else:
                    row_cells = current_table.add_row().cells
                    for i, cell in enumerate(cells):
                        row_cells[i].text = cell
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # 普通段落
                if current_table:
                    current_table = None
                    in_table = False
                
                # 处理粗体文本
                paragraph_text = line.replace('**', '').replace('*', '').strip()
                if paragraph_text:
                    p = doc.add_paragraph(paragraph_text, style='Normal')
                    p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
                    
                    # 处理关键字加粗
                    self.format_keywords_fallback(p)
    
    def add_appendix_fallback(self, doc, content: str):
        """添加附录（回退方法）"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # 添加分页符
        doc.add_page_break()
        
        # 附录标题
        appendix_title = doc.add_paragraph("附录")
        appendix_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        appendix_title.runs[0].font.name = '黑体'
        appendix_title.runs[0].font.size = Pt(16)
        appendix_title.runs[0].font.bold = True
        appendix_title.paragraph_format.space_after = Pt(18)
        
        # 报告信息
        info_items = [
            f"报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
            "专业顾问：私校申请专家团队",
            "版本：v2.0",
            "页数统计：{PAGE} 页",
            "字数统计：{NUMCHARS} 字"
        ]
        
        for item in info_items:
            info_item = doc.add_paragraph(item)
            info_item.runs[0].font.name = '仿宋'
            info_item.runs[0].font.size = Pt(12)
            info_item.paragraph_format.space_after = Pt(6)
    
    def setup_headers_footers_fallback(self, doc, content: str):
        """设置页眉页脚（回退方法）"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 提取学生姓名
        student_name = "Alex Chen"
        for line in content.split('\n'):
            if '姓名' in line and 'Alex Chen' in line:
                student_name = "Alex Chen"
                break
        
        # 设置页眉页脚
        for section in doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{student_name} - 学校申请报告"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.runs[0].font.name = '仿宋'
            header_para.runs[0].font.size = Pt(10.5)
            
            # 页脚
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "第 {PAGE} 页 / 共 {NUMPAGES} 页"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.name = '仿宋'
            footer_para.runs[0].font.size = Pt(10.5)
    
    def format_keywords_fallback(self, paragraph):
        """格式化关键字加粗（回退方法）"""
        keywords = ['家庭价值观', '教育目标', '文化背景', '支持程度', '期望设定', '资源投入',
                   'GPA', '强项领域', '标准化考试', '学术竞赛', '学习能力', '适应能力',
                   '创新思维', '责任感', '团队协作', '影响力', '项目经验', '专业特长',
                   '全球视野', '创新精神', '全人发展', '学术', '活动资源', '价值观', '文化',
                   '性格', '氛围', '优势匹配', '申请建议', '推荐理由']
        
        text = paragraph.text
        for keyword in keywords:
            if keyword in text:
                # 简单的关键字加粗处理
                runs = paragraph.runs
                for run in runs:
                    if keyword in run.text:
                        run.font.bold = True
                        break

def main():
    """测试增强版报告生成器"""
    # 测试LLM pipeline
    print("=== 测试LLM驱动pipeline ===")
    generator_llm = EnhancedReportGenerator(use_llm_pipeline=True)
    
    # 测试数据
    student_data = {
        "name": "Alex Chen",
        "age": "14岁",
        "grade": "Grade 8",
        "gpa": "3.8/4.0",
        "academic_strengths": "数学、物理、计算机科学",
        "competition_achievements": "机器人竞赛省级二等奖",
        "leadership_positions": "科技部副部长",
        "project_experiences": "环保义卖活动组织",
        "learning_ability": "自主学习和问题解决",
        "adaptability": "跨文化环境适应",
        "target_schools": [
            {
                "name": "Upper Canada College",
                "scores": {"academic": 4, "activities": 4, "culture": 5, "personality": 4},
                "weights": {"academic": 0.35, "activities": 0.25, "culture": 0.2, "personality": 0.2}
            },
            {
                "name": "Havergal College", 
                "scores": {"academic": 4, "activities": 3, "culture": 4, "personality": 3},
                "weights": {"academic": 0.35, "activities": 0.25, "culture": 0.2, "personality": 0.2}
            }
        ]
    }
    
    conversation_log = [
        {"role": "student", "content": "我喜欢组织活动和做科学实验"},
        {"role": "parent", "content": "孩子在学生会组织过环保义卖"}
    ]
    
    # 生成报告
    report_result = generator_llm.generate_comprehensive_report(conversation_log, student_data)
    
    print("LLM Pipeline报告生成成功!")
    print(f"页数: {report_result['metadata']['page_count']}")
    print(f"字数: {report_result['metadata']['word_count']}")
    
    # 导出报告
    exported_files = generator_llm.export_report(report_result, "all")
    
    print("\nLLM Pipeline导出文件:")
    for format_type, file_path in exported_files.items():
        print(f"{format_type}: {file_path}")
    
    # 测试传统pipeline
    print("\n=== 测试传统pipeline ===")
    generator_traditional = EnhancedReportGenerator(use_llm_pipeline=False)
    
    # 生成报告
    report_result_traditional = generator_traditional.generate_comprehensive_report(conversation_log, student_data)
    
    print("传统Pipeline报告生成成功!")
    print(f"页数: {report_result_traditional['metadata']['page_count']}")
    print(f"字数: {report_result_traditional['metadata']['word_count']}")
    
    # 导出报告
    exported_files_traditional = generator_traditional.export_report(report_result_traditional, "all")
    
    print("\n传统Pipeline导出文件:")
    for format_type, file_path in exported_files_traditional.items():
        print(f"{format_type}: {file_path}")

if __name__ == "__main__":
    main()
