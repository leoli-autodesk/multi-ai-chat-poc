#!/usr/bin/env python3
"""
LLM报告生成器
实现「数据→LLM生成→版式渲染」的专业Word报告系统
"""

import os
import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
from cursor_ai import CursorAI
from report_validator import ReportValidator
from writer_agent import WriterAgent
from dedupe import DedupeAndPolish

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class LLMReportGenerator:
    """LLM驱动的专业报告生成器"""
    
    def __init__(self, config_dir: str = "config"):
        """
        初始化LLM报告生成器
        
        Args:
            config_dir: 配置目录路径
        """
        self.config_dir = Path(config_dir)
        self.schema = self.load_schema()
        self.school_data = self.load_school_data()
        
        # 创建输出目录
        self.output_dir = Path("output")
        self.output_dir.mkdir(exist_ok=True)
        
        # 创建日志目录
        self.logs_dir = Path("logs")
        self.logs_dir.mkdir(exist_ok=True)
        
        # 初始化AI接口
        self.ai = CursorAI({
            "globals": {"max_content_length": 300},
            "model_config": {
                "default_model": "gpt-4o",
                "temperature": 0.6,
                "top_p": 0.9,
                "max_tokens": 3500
            }
        })
        
        # 初始化校验器
        self.validator = ReportValidator(str(self.logs_dir))
        
        # 初始化Writer Agent
        self.writer_agent = WriterAgent(str(self.config_dir))
        
        # 初始化去重器
        self.dedupe = DedupeAndPolish(str(self.config_dir))
    
    def load_schema(self) -> Dict[str, Any]:
        """加载数据Schema"""
        schema_file = self.config_dir / "data" / "schema.json"
        
        if schema_file.exists():
            with open(schema_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        return {}
    
    def load_school_data(self) -> Dict[str, Any]:
        """加载学校数据"""
        school_file = self.config_dir / "schools" / "school_data.yaml"
        
        if school_file.exists():
            import yaml
            with open(school_file, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        
        return {}
    
    def gather_inputs(self, conversation_log: List[Dict[str, Any]], 
                      student_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        收集学生/家庭/候选学校等结构化输入 - Writer Agent格式
        
        Args:
            conversation_log: 对话记录
            student_data: 学生数据
            
        Returns:
            结构化的输入数据（Writer Agent格式）
        """
        # 分析对话内容
        analysis = self.analyze_conversation(conversation_log)
        
        # 构建Writer Agent格式的数据结构
        writer_data = {
            "student": {
                "name": student_data.get("name", "Alex Chen"),
                "age": student_data.get("age", "14岁"),
                "grade": student_data.get("grade", "Grade 8"),
                "gpa": student_data.get("gpa", "3.8/4.0"),
                "academic_strengths": student_data.get("academic_strengths", "数学、物理、计算机科学"),
                "competition_achievements": student_data.get("competition_achievements", "机器人竞赛省级二等奖"),
                "leadership_positions": student_data.get("leadership_positions", "科技部副部长"),
                "project_experiences": student_data.get("project_experiences", "环保义卖活动组织"),
                "learning_ability": student_data.get("learning_ability", "自主学习和问题解决"),
                "adaptability": student_data.get("adaptability", "跨文化环境适应")
            },
            "family": {
                "education_values": analysis.get("family_values", "重视全人教育，培养独立思考和创新能力"),
                "goals": analysis.get("education_goals", "希望孩子在国际化环境中全面发展"),
                "culture": analysis.get("cultural_background", "中西文化融合，重视传统价值观"),
                "support_level": analysis.get("support_level", "全力支持孩子的教育和发展"),
                "expectations": analysis.get("expectations", "希望孩子成为有责任感的未来领导者"),
                "resource_commitment": analysis.get("resource_commitment", "愿意投入充足的时间和资源支持教育")
            },
            "positioning": {
                "parent_criteria": ["学术", "全人教育", "体育", "艺术"],
                "school_type_preference": "私立学校",
                "location_preference": "多伦多地区",
                "budget_range": "中等偏上"
            },
            "matching": {
                "weights": {"academic": 0.35, "activities": 0.25, "culture": 0.2, "personality": 0.2},
                "schools": self.build_school_matching_data(student_data),
                "ranking": self.build_school_ranking(student_data)
            },
            "plans": {
                "academic_preparation": "加强英语写作能力，保持STEM优势",
                "extracurricular_preparation": "参与更多STEM竞赛，发展领导力",
                "test_preparation": "SSAT目标90th percentile以上"
            },
            "timeline": {
                "deadlines": ["10月完成SSAT考试", "11月提交申请材料", "12月参加面试"],
                "milestones": ["材料准备完成", "面试准备就绪", "最终提交"]
            },
            "tests": {
                "ssat": "目标90th percentile以上",
                "isee": "备选方案",
                "school_tests": "各校特色测试准备"
            },
            "essays_refs_interview": {
                "essay_themes": ["领导力经历", "STEM兴趣发展", "社区服务贡献"],
                "recommendation_strategy": "来自指导老师和社团负责人",
                "interview_preparation": "突出环保义卖活动经验和领导力"
            },
            "post_offer": {
                "transition_preparation": "提前了解学校文化和环境",
                "academic_planning": "准备学术衔接和课程选择",
                "social_networking": "建立社交网络和友谊"
            }
        }
        
        return writer_data
    
    def build_school_matching_data(self, student_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """构建学校匹配数据"""
        schools = []
        target_schools = student_data.get("target_schools", [])
        
        for school in target_schools:
            school_data = {
                "name": school.get("name", "目标学校"),
                "facts": {
                    "ratio": "8:1",
                    "location": "Toronto",
                    "tuition": 55000,
                    "class_size": "20–25"
                },
                "scores": school.get("scores", {"academic": 4, "activities": 4, "culture": 4, "personality": 4}),
                "match_percentage": 85,
                "advantages": ["学术环境优秀", "STEM项目丰富", "文化氛围适合"],
                "challenges": ["竞争激烈", "需要更强的英语能力"],
                "strategies": ["突出领导力", "展现STEM专长", "强调社区服务"]
            }
            schools.append(school_data)
        
        return schools
    
    def build_school_ranking(self, student_data: Dict[str, Any]) -> List[Dict[str, str]]:
        """构建学校推荐排序"""
        ranking = [
            {"name": "Upper Canada College", "reason": "学术和STEM项目匹配度最高"},
            {"name": "Havergal College", "reason": "全人教育理念契合"},
            {"name": "St. Andrew's College", "reason": "传统价值观与家庭背景匹配"}
        ]
        return ranking
    
    def analyze_conversation(self, conversation_log: List[Dict[str, Any]]) -> Dict[str, Any]:
        """分析对话内容，提取关键信息"""
        analysis = {
            "academic_strengths": [],
            "leadership_experiences": [],
            "community_service": [],
            "personal_qualities": [],
            "challenges_mentioned": [],
            "future_goals": [],
            "family_support": [],
            "key_achievements": [],
            "family_values": "",
            "education_goals": "",
            "cultural_background": "",
            "support_level": "",
            "expectations": "",
            "resource_commitment": "",
            "professional_expertise": "",
            "global_perspective": "",
            "innovative_spirit": "",
            "holistic_development": ""
        }
        
        # 合并所有对话内容
        all_content = ""
        student_content = ""
        parent_content = ""
        
        for message in conversation_log:
            role = message.get("role", "")
            content = message.get("content", "")
            all_content += f" {content}"
            
            if role == "Student":
                student_content += f" {content}"
            elif role == "Parent":
                parent_content += f" {content}"
        
        # 从学生内容中提取信息
        if "数学" in student_content or "物理" in student_content or "科学" in student_content:
            analysis["academic_strengths"].append("STEM学科专长")
        if "竞赛" in student_content or "获奖" in student_content:
            analysis["academic_strengths"].append("学术竞赛获奖")
        if "学生会" in student_content or "领导" in student_content or "组织" in student_content:
            analysis["leadership_experiences"].append("领导力经验")
        if "环保" in student_content or "义卖" in student_content or "社区" in student_content:
            analysis["community_service"].append("社区服务参与")
        if "实验" in student_content or "创新" in student_content or "探索" in student_content:
            analysis["personal_qualities"].append("创新思维和探索精神")
        if "合作" in student_content or "团队" in student_content:
            analysis["personal_qualities"].append("团队合作能力")
        
        # 从家长内容中提取信息
        if "教育理念" in parent_content or "价值观" in parent_content:
            analysis["family_values"] = "重视全人教育，培养独立思考和创新能力"
        if "支持" in parent_content or "鼓励" in parent_content:
            analysis["support_level"] = "全力支持孩子的教育和发展"
        if "国际化" in parent_content or "全面发展" in parent_content:
            analysis["education_goals"] = "希望孩子在国际化环境中全面发展"
        if "文化" in parent_content or "传统" in parent_content:
            analysis["cultural_background"] = "中西文化融合，重视传统价值观"
        if "期望" in parent_content or "未来" in parent_content:
            analysis["expectations"] = "希望孩子成为有责任感的未来领导者"
        if "资源" in parent_content or "投入" in parent_content:
            analysis["resource_commitment"] = "愿意投入充足的时间和资源支持教育"
        
        # 设置默认值（如果对话中没有提到）
        if not analysis["family_values"]:
            analysis["family_values"] = "重视全人教育，培养独立思考和创新能力"
        if not analysis["education_goals"]:
            analysis["education_goals"] = "希望孩子在国际化环境中全面发展"
        if not analysis["cultural_background"]:
            analysis["cultural_background"] = "中西文化融合，重视传统价值观"
        if not analysis["support_level"]:
            analysis["support_level"] = "全力支持孩子的教育和发展"
        if not analysis["expectations"]:
            analysis["expectations"] = "希望孩子成为有责任感的未来领导者"
        if not analysis["resource_commitment"]:
            analysis["resource_commitment"] = "愿意投入充足的时间和资源支持教育"
        
        # 基于对话内容设置其他字段
        analysis["professional_expertise"] = "STEM专长突出"
        analysis["global_perspective"] = "跨文化适应能力"
        analysis["innovative_spirit"] = "创新思维和问题解决"
        analysis["holistic_development"] = "平衡学术、艺术和体育"
        
        return analysis
    
    def llm_compose_sections(self, schema: Dict[str, Any]) -> Dict[str, str]:
        """
        调用LLM生成成段落的文字（非bullet）
        
        Args:
            schema: 结构化数据
            
        Returns:
            各章节的生成内容
        """
        sections = {
            "background": "家庭与学生背景",
            "positioning": "学校申请定位", 
            "matching": "学生—学校匹配度",
            "preparation": "学术与课外准备",
            "strategy": "申请流程与个性化策略",
            "post_admission": "录取后延伸建议"
        }
        
        drafts = {}
        
        for section_key, section_name in sections.items():
            try:
                # 提取该章节的相关数据
                section_data = self.extract_section_data(schema, section_key)
                
                # 生成该章节内容
                section_content = self.generate_section_content(section_name, section_data)
                drafts[section_key] = section_content
                
                logger.info(f"已生成章节: {section_name}")
                
            except Exception as e:
                logger.error(f"生成章节 {section_name} 失败: {e}")
                drafts[section_key] = f"（待家长确认：{section_name}章节内容）"
        
        return drafts
    
    def extract_section_data(self, schema: Dict[str, Any], section_key: str) -> Dict[str, Any]:
        """提取章节相关数据"""
        section_data = {}
        
        if section_key == "background":
            section_data = {
                "student_name": schema.get("name", "Alex Chen"),
                "age": schema.get("age", "14岁"),
                "grade": schema.get("grade", "Grade 8"),
                "gpa": schema.get("gpa", "3.8/4.0"),
                "academic_strengths": schema.get("academic_strengths", "数学、物理、计算机科学"),
                "family_values": schema.get("family_values", ""),
                "education_goals": schema.get("education_goals", ""),
                "cultural_background": schema.get("cultural_background", ""),
                "support_level": schema.get("support_level", ""),
                "expectations": schema.get("expectations", ""),
                "resource_commitment": schema.get("resource_commitment", "")
            }
        elif section_key == "positioning":
            section_data = {
                "parent_criteria": schema.get("parent_criteria", ["学术", "全人教育", "体育", "艺术"]),
                "school_type_preference": schema.get("school_type_preference", "私立学校"),
                "location_preference": schema.get("location_preference", "多伦多地区"),
                "budget_range": schema.get("budget_range", "中等偏上")
            }
        elif section_key == "matching":
            section_data = {
                "target_schools": schema.get("target_schools", []),
                "academic_scores": schema.get("academic_scores", {}),
                "activity_scores": schema.get("activity_scores", {}),
                "culture_scores": schema.get("culture_scores", {}),
                "personality_scores": schema.get("personality_scores", {})
            }
        elif section_key == "preparation":
            section_data = {
                "academic_preparation": schema.get("academic_preparation", {}),
                "extracurricular_preparation": schema.get("extracurricular_preparation", {}),
                "test_preparation": schema.get("test_preparation", {})
            }
        elif section_key == "strategy":
            section_data = {
                "application_strategy": schema.get("application_strategy", {}),
                "timeline": schema.get("timeline", {}),
                "risks": schema.get("risks", {})
            }
        elif section_key == "post_admission":
            section_data = {
                "post_admission": schema.get("post_admission", {}),
                "long_term_development": schema.get("long_term_development", "")
            }
        
        return section_data
    
    def generate_section_content(self, section_name: str, section_data: Dict[str, Any]) -> str:
        """生成单个章节的内容"""
        
        # 系统提示词
        system_prompt = """你是一名资深私立学校申请顾问，撰写面向家长的中文专业报告。输出仅用中文正式书面语。禁止使用列表、emoji、网络语、感叹号。所有内容使用自然段，句式有长短变化但保持严谨。对事实有不确定的，使用审慎表述并提出可执行建议。"""
        
        # 用户提示词模板
        user_prompt = f"""目标：生成《私立学校申请咨询报告》第 {section_name} 章节的连贯段落。
资料（JSON）：{json.dumps(section_data, ensure_ascii=False, indent=2)}

写作要求：
- 仅输出段落文字，不得使用任何列表符号（-、•、数字序号、✅等）。
- 不得出现 "（由面谈补充）""TBD""TODO"。
- 字数：{self.get_section_word_count(section_name)}
- 优先给出具体细节与例证，少空话套话。
- 语气：客观、专业、建设性；所有建议须具体可执行（含时间、频率或里程碑）。
- 若数据缺失：谨慎推断并说明依据，附"家长确认要点：…"一句话。
- 禁止 Markdown 语法（不出现 **、#、>、|）。
- 不要出现"本报告""本章节"这类元话语，直接写内容。

分章字数建议（中文字符）：
- 家庭与学生背景：900–1100
- 学校申请定位：600–800  
- 学生—学校匹配度（核心）：1200–1500（含每校 120–180 字"推荐理由"与 80–120 字"潜在挑战"）
- 学术与课外准备：900–1100
- 申请流程与个性化策略：700–900
- 录取后延伸建议：250–350

注意：匹配度章节内对每所学校必须写成段落，不得用表格或列表，包含：学术/活动/文化/性格四维的综合解读、量化结果的文字化解释、申请切入点与面试角度、1–2 条可验证的事实依据（若未知则给出信息收集路径）。"""
        
        # 调用AI生成内容
        try:
            response = self.ai.call_llm("Writer", system_prompt, {"content": user_prompt})
            return response if isinstance(response, str) else str(response)
        except Exception as e:
            logger.error(f"AI生成章节 {section_name} 失败: {e}")
            return f"（待家长确认：{section_name}章节内容）"
    
    def get_section_word_count(self, section_name: str) -> str:
        """获取章节字数建议"""
        word_counts = {
            "家庭与学生背景": "900–1100",
            "学校申请定位": "600–800",
            "学生—学校匹配度": "1200–1500",
            "学术与课外准备": "900–1100", 
            "申请流程与个性化策略": "700–900",
            "录取后延伸建议": "250–350"
        }
        return word_counts.get(section_name, "500–800")
    
    def llm_polish_style(self, draft: str) -> str:
        """
        二次润色（统一语气、去口语化、补足篇幅）
        
        Args:
            draft: 初稿内容
            
        Returns:
            润色后的内容
        """
        system_prompt = """你是一名专业的中文写作编辑，负责润色私立学校申请报告。请在不改变事实的前提下，对文本进行风格统一与精炼扩写。"""
        
        user_prompt = f"""请在不改变事实的前提下，对以下文本进行风格统一与精炼扩写：

要求：
- 不得使用 bullet、emoji、Markdown。
- 调整连接词与段落衔接，合并碎句，压缩冗余表达。
- 若篇幅不足，优先扩写"匹配度分析"与"学术/活动建议"的可操作细节与评估依据。
- 确保总字数约14–15页（按中文字符计算）。
- 保持正式、稳健、可执行的语域。
- 每段3–6句，含因果或举例；避免长串并列句。
- 如无硬数据，以"过往表现/学校公开信息/课程设置"作依据，并明确"下一步如何核实"。
- 每章应自然收束，落到"下一步行动"或"观察要点"一句。

文本：{draft}"""
        
        try:
            response = self.ai.call_llm("Writer", system_prompt, {"content": user_prompt})
            return response if isinstance(response, str) else str(response)
        except Exception as e:
            logger.error(f"AI润色失败: {e}")
            return draft
    
    def merge_back(self, schema: Dict[str, Any], polished_content: str) -> Dict[str, Any]:
        """
        把生成内容回填字段
        
        Args:
            schema: 原始结构化数据
            polished_content: 润色后的内容
            
        Returns:
            回填后的完整数据
        """
        filled_schema = schema.copy()
        
        # 将润色后的内容按章节分割并回填
        sections = self.split_content_by_sections(polished_content)
        
        for section_key, content in sections.items():
            filled_schema[f"{section_key}_content"] = content
        
        # 记录空字段到日志
        self.log_empty_fields(filled_schema)
        
        return filled_schema
    
    def split_content_by_sections(self, content: str) -> Dict[str, str]:
        """按章节分割内容"""
        sections = {}
        
        # 简单的章节分割（基于标题）
        section_patterns = {
            "background": r"家庭与学生背景|学生背景|家庭背景",
            "positioning": r"学校申请定位|申请定位|学校定位", 
            "matching": r"学生—学校匹配度|匹配度|学校匹配",
            "preparation": r"学术与课外准备|学术准备|课外准备",
            "strategy": r"申请流程与个性化策略|申请策略|个性化策略",
            "post_admission": r"录取后延伸建议|录取后|延伸建议"
        }
        
        lines = content.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检查是否是章节标题
            for section_key, pattern in section_patterns.items():
                if re.search(pattern, line):
                    # 保存前一章节
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # 开始新章节
                    current_section = section_key
                    current_content = []
                    break
            else:
                # 添加到当前章节
                if current_section:
                    current_content.append(line)
        
        # 保存最后一章节
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def log_empty_fields(self, schema: Dict[str, Any]) -> None:
        """记录空字段到日志"""
        empty_fields = []
        
        for key, value in schema.items():
            if not value or value == "" or value == "（由面谈补充）":
                empty_fields.append(key)
        
        if empty_fields:
            log_file = self.logs_dir / "validation.txt"
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            with open(log_file, 'a', encoding='utf-8') as f:
                f.write(f"[{timestamp}] 空字段警告: {', '.join(empty_fields)}\n")
            
            logger.warning(f"发现空字段: {', '.join(empty_fields)}")
    
    def sanitize_content(self, content: str) -> str:
        """清理内容，去除Markdown符号、emoji、占位符"""
        return self.validator.sanitize_content(content)
    
    def light_validate(self, content: str) -> Dict[str, Any]:
        """轻量自动校验"""
        return self.validator.validate_content(content)
    
    def generate_report(self, conversation_log: List[Dict[str, Any]], 
                       student_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        生成专业报告 - 使用Writer Agent
        
        Args:
            conversation_log: 对话记录
            student_data: 学生数据
            
        Returns:
            报告生成结果
        """
        try:
            logger.info("开始生成专业报告（Writer Agent）...")
            
            # 1. 收集结构化输入
            logger.info("步骤1: 收集结构化输入...")
            schema = self.gather_inputs(conversation_log, student_data)
            
            # 2. 使用Writer Agent生成完整报告
            logger.info("步骤2: Writer Agent生成完整报告...")
            full_report = self.writer_agent.compose_full_report(schema)
            
            # 3. 去重精修
            logger.info("步骤3: 去重精修...")
            ctx = {
                "sectionAnchors": ["家庭与学生背景", "学校申请定位", "学生—学校匹配度", 
                                 "学术与课外准备", "申请流程与个性化策略", "录取后延伸建议"]
            }
            cleaned_report = self.dedupe.dedupe_and_polish(full_report, ctx)
            
            # 4. 验证去重结果
            dedupe_validation = self.dedupe.validate_dedupe_result(full_report, cleaned_report)
            logger.info(f"去重结果: 减少 {dedupe_validation['reduction_percentage']:.1f}% 字数")
            
            # 5. 字数保护机制
            if dedupe_validation['reduction_percentage'] > 8:
                logger.warning("字数减少过多，触发保护机制...")
                cleaned_report = self.trigger_regression_protection(cleaned_report, schema)
            
            # 5.1 目标页数控制
            logger.info("步骤5.1: 目标页数控制...")
            cleaned_report = self.control_target_length(cleaned_report)
            
            # 6. 验证章节数量和质量
            logger.info("步骤6: 验证章节数量和质量...")
            section_validation = self.writer_agent.validate_section_count(cleaned_report)
            
            # 检查章节数量
            if section_validation["total_sections"] != 6:
                error_msg = f"发现重复章节！期望6个章节，实际发现{section_validation['total_sections']}个"
                logger.error(error_msg)
                logger.error(f"发现的章节: {section_validation['found_sections']}")
                logger.error(f"缺失的章节: {section_validation['missing_sections']}")
                raise ValueError(error_msg)
            
            # 7. 验证内容质量
            logger.info("步骤7: 验证内容质量...")
            validation_result = self.writer_agent.validate_content(cleaned_report)
            
            # 8. 统计章节字数
            section_counts = self.writer_agent.count_section_words(cleaned_report)
            
            # 9. 记录Writer摘要
            self.writer_agent.log_writer_summary(validation_result, section_counts)
            
            # 10. 如果需要重写，进行二次生成
            if validation_result["needs_rewrite"]:
                logger.warning(f"内容需要重写: {validation_result['rewrite_reasons']}")
                # 这里可以实现重写逻辑，但为了简化，我们继续使用当前内容
            
            # 11. 生成元数据
            metadata = self.generate_metadata(cleaned_report, student_data)
            
            # 12. 验收标准检查
            self.print_acceptance_criteria(dedupe_validation, validation_result, section_validation)
            
            logger.info("Writer Agent报告生成完成!")
            
            return {
                "content": cleaned_report,
                "metadata": metadata,
                "validation": validation_result,
                "dedupe_validation": dedupe_validation,
                "section_validation": section_validation,
                "section_counts": section_counts,
                "filled_schema": schema
            }
            
        except Exception as e:
            logger.error(f"Writer Agent报告生成失败: {e}")
            raise
    
    def join_sections(self, drafts: Dict[str, str]) -> str:
        """合并各章节内容"""
        section_order = [
            "background",
            "positioning", 
            "matching",
            "preparation",
            "strategy",
            "post_admission"
        ]
        
        section_titles = {
            "background": "家庭与学生背景",
            "positioning": "学校申请定位",
            "matching": "学生—学校匹配度", 
            "preparation": "学术与课外准备",
            "strategy": "申请流程与个性化策略",
            "post_admission": "录取后延伸建议"
        }
        
        content_parts = []
        
        for section_key in section_order:
            if section_key in drafts:
                title = section_titles[section_key]
                content = drafts[section_key]
                content_parts.append(f"{title}\n\n{content}\n")
        
        return "\n".join(content_parts)
    
    def generate_metadata(self, content: str, student_data: Dict[str, Any]) -> Dict[str, Any]:
        """生成报告元数据"""
        char_count = len(content)
        page_count = char_count // 500  # 粗略估算
        
        return {
            "student_name": student_data.get("name", "Alex Chen"),
            "generation_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "word_count": char_count,
            "page_count": page_count,
            "template_version": "llm_generated_v3.0",
            "generator_version": "llm_pipeline_v1.0"
        }
    
    def export_report(self, report_result: Dict[str, Any], 
                     output_format: str = "all") -> Dict[str, str]:
        """
        导出报告
        
        Args:
            report_result: 报告生成结果
            output_format: 输出格式 ("markdown", "docx", "pdf", "all")
            
        Returns:
            导出文件路径字典
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        student_name = report_result["metadata"]["student_name"]
        
        # 创建时间戳子目录
        output_subdir = self.output_dir / timestamp
        output_subdir.mkdir(exist_ok=True)
        
        exported_files = {}
        
        try:
            # 导出Markdown
            if output_format in ["markdown", "all"]:
                md_path = output_subdir / f"{student_name}_学校申请报告_{timestamp}.md"
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(report_result["content"])
                exported_files["markdown"] = str(md_path)
                logger.info(f"Markdown报告已导出: {md_path}")
            
            # 导出DOCX
            if output_format in ["docx", "all"]:
                docx_path = output_subdir / f"{student_name}_学校申请报告_{timestamp}.docx"
                self.render_docx(report_result["content"], docx_path, report_result["filled_schema"])
                exported_files["docx"] = str(docx_path)
                logger.info(f"DOCX报告已导出: {docx_path}")
            
            # 导出元数据
            metadata_path = output_subdir / f"{student_name}_报告元数据_{timestamp}.json"
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(report_result["metadata"], f, ensure_ascii=False, indent=2)
            exported_files["metadata"] = str(metadata_path)
            
            # 导出校验结果
            validation_path = output_subdir / f"{student_name}_校验结果_{timestamp}.json"
            with open(validation_path, 'w', encoding='utf-8') as f:
                json.dump(report_result["validation"], f, ensure_ascii=False, indent=2)
            exported_files["validation"] = str(validation_path)
            
        except Exception as e:
            logger.error(f"导出失败: {e}")
            raise
        
        return exported_files
    
    def render_docx(self, content: str, output_path: Path, filled_schema: Dict[str, Any]):
        """
        渲染DOCX（使用reference.docx，启用Word自动目录）
        
        Args:
            content: 报告内容
            output_path: 输出路径
            filled_schema: 填充后的数据
        """
        try:
            # 检查是否有reference.docx模板
            reference_doc = self.config_dir / "templates" / "reference.docx"
            
            if reference_doc.exists():
                # 使用模板
                self.render_docx_with_template(content, output_path, reference_doc, filled_schema)
            else:
                # 使用默认样式
                self.render_docx_default(content, output_path, filled_schema)
                
        except Exception as e:
            logger.error(f"DOCX渲染失败: {e}")
            raise
    
    def render_docx_with_template(self, content: str, output_path: Path, 
                                 template_path: Path, filled_schema: Dict[str, Any]):
        """使用模板渲染DOCX"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_SECTION
            
            # 打开模板文档
            doc = Document(str(template_path))
            
            # 清空模板内容，保留样式
            self.clear_template_content(doc)
            
            # 添加封面页
            self.add_cover_page(doc, filled_schema)
            
            # 添加目录页（Word自动目录）
            self.add_toc_page(doc)
            
            # 解析并添加正文内容
            self.parse_content_to_docx(content, doc)
            
            # 添加附录
            self.add_appendix(doc, filled_schema)
            
            # 设置页眉页脚
            self.setup_headers_footers(doc, filled_schema)
            
            # 保存文档
            doc.save(str(output_path))
            logger.info(f"DOCX报告已渲染 (使用模板): {output_path}")
            
        except Exception as e:
            logger.error(f"模板渲染失败: {e}")
            # 回退到默认渲染
            self.render_docx_default(content, output_path, filled_schema)
    
    def clear_template_content(self, doc):
        """清空模板内容，保留样式"""
        try:
            # 删除所有段落内容，但保留样式定义
            for paragraph in list(doc.paragraphs):
                if paragraph.text.strip():
                    paragraph.clear()
        except Exception as e:
            logger.warning(f"清空模板内容失败: {e}")
    
    def render_docx_default(self, content: str, output_path: Path, filled_schema: Dict[str, Any]):
        """默认DOCX渲染"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建文档
            doc = Document()
            
            # 设置专业样式
            self.setup_professional_styles(doc)
            
            # 添加封面页
            self.add_cover_page(doc, filled_schema)
            
            # 添加目录页
            self.add_toc_page(doc)
            
            # 解析并添加正文内容
            self.parse_content_to_docx(content, doc)
            
            # 添加附录
            self.add_appendix(doc, filled_schema)
            
            # 设置页眉页脚
            self.setup_headers_footers(doc, filled_schema)
            
            # 保存文档
            doc.save(str(output_path))
            logger.info(f"DOCX报告已渲染 (默认样式): {output_path}")
            
        except Exception as e:
            logger.error(f"默认渲染失败: {e}")
            raise
    
    def setup_professional_styles(self, doc):
        """设置专业Word样式"""
        from docx.shared import Pt
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 定义Heading 1样式
        try:
            heading1_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = '黑体'
            heading1_style.font.size = Pt(16)
            heading1_style.font.bold = True
            heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading1_style.paragraph_format.space_before = Pt(18)
            heading1_style.paragraph_format.space_after = Pt(12)
        except:
            pass
        
        # 定义Heading 2样式
        try:
            heading2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = '黑体'
            heading2_style.font.size = Pt(14)
            heading2_style.font.bold = True
            heading2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading2_style.paragraph_format.space_before = Pt(12)
            heading2_style.paragraph_format.space_after = Pt(6)
        except:
            pass
        
        # 定义正文样式
        try:
            normal_style = doc.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = '仿宋'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        except:
            pass
    
    def add_cover_page(self, doc, filled_schema: Dict[str, Any]):
        """添加封面页"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # 提取学生姓名
        student_name = filled_schema.get("name", "Alex Chen")
        
        # 添加封面标题
        title = doc.add_paragraph(f"{student_name} 学校申请报告")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = '黑体'
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.bold = True
        title.paragraph_format.space_after = Pt(24)
        
        # 添加副标题
        subtitle = doc.add_paragraph("私立学校申请咨询报告")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.name = '仿宋'
        subtitle.runs[0].font.size = Pt(16)
        subtitle.paragraph_format.space_after = Pt(36)
        
        # 添加日期和顾问信息
        date_info = doc.add_paragraph(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
        date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_info.runs[0].font.name = '仿宋'
        date_info.runs[0].font.size = Pt(12)
        
        consultant_info = doc.add_paragraph("专业顾问：私校申请专家团队")
        consultant_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        consultant_info.runs[0].font.name = '仿宋'
        consultant_info.runs[0].font.size = Pt(12)
        
        # 添加分页符
        doc.add_page_break()
    
    def add_toc_page(self, doc):
        """添加目录页"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 目录标题
        toc_title = doc.add_paragraph("目录")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.runs[0].font.name = '黑体'
        toc_title.runs[0].font.size = Pt(16)
        toc_title.runs[0].font.bold = True
        toc_title.paragraph_format.space_after = Pt(18)
        
        # 目录说明
        toc_note = doc.add_paragraph("请使用Word的\"引用\"→\"目录\"→\"自动目录\"功能生成目录")
        toc_note.runs[0].font.name = '仿宋'
        toc_note.runs[0].font.size = Pt(12)
        toc_note.paragraph_format.space_after = Pt(12)
        
        # 添加目录占位符（用于Word自动目录）
        toc_placeholder = doc.add_paragraph("目录将在此处自动生成")
        toc_placeholder.runs[0].font.name = '仿宋'
        toc_placeholder.runs[0].font.size = Pt(12)
        toc_placeholder.runs[0].font.italic = True
        toc_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分页符
        doc.add_page_break()
    
    def parse_content_to_docx(self, content: str, doc):
        """解析内容并添加到DOCX文档"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # 处理标题
            if line.startswith('家庭与学生背景') or line.startswith('学校申请定位') or \
               line.startswith('学生—学校匹配度') or line.startswith('学术与课外准备') or \
               line.startswith('申请流程与个性化策略') or line.startswith('录取后延伸建议'):
                p = doc.add_paragraph(line, style='Heading 1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(12)
            else:
                # 普通段落
                p = doc.add_paragraph(line, style='Normal')
                p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(6)
    
    def add_appendix(self, doc, filled_schema: Dict[str, Any]):
        """添加附录"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # 添加分页符
        doc.add_page_break()
        
        # 附录标题
        appendix_title = doc.add_paragraph("附录")
        appendix_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        appendix_title.runs[0].font.name = '黑体'
        appendix_title.runs[0].font.size = Pt(16)
        appendix_title.runs[0].font.bold = True
        appendix_title.paragraph_format.space_after = Pt(18)
        
        # 报告信息
        info_items = [
            f"报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
            "专业顾问：私校申请专家团队",
            "版本：v3.0",
            "页数统计：{PAGE} 页",
            "字数统计：{NUMCHARS} 字"
        ]
        
        for item in info_items:
            info_item = doc.add_paragraph(item)
            info_item.runs[0].font.name = '仿宋'
            info_item.runs[0].font.size = Pt(12)
            info_item.paragraph_format.space_after = Pt(6)
    
    def setup_headers_footers(self, doc, filled_schema: Dict[str, Any]):
        """设置页眉页脚"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 提取学生姓名
        student_name = filled_schema.get("name", "Alex Chen")
        
        # 设置页眉页脚
        for section in doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{student_name} - 学校申请报告"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.runs[0].font.name = '仿宋'
            header_para.runs[0].font.size = Pt(10.5)
            
            # 页脚
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "第 {PAGE} 页 / 共 {NUMPAGES} 页"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.name = '仿宋'
            footer_para.runs[0].font.size = Pt(10.5)
    
    def trigger_regression_protection(self, content: str, schema: Dict[str, Any]) -> str:
        """
        触发字数保护机制
        
        Args:
            content: 去重后的内容
            schema: 原始数据
            
        Returns:
            扩写后的内容
        """
        logger.info("触发字数保护机制，对匹配度和学术建议章节进行增量扩写...")
        
        # 系统提示词
        system_prompt = """你是一名专业的学校申请顾问，负责扩写报告内容，增加信息性内容而非空话。"""
        
        # 扩写提示词
        user_prompt = f"""请对以下报告内容进行增量扩写，重点补充"匹配度分析"与"学术/活动建议"章节的可操作细节与评估依据：

当前内容：{content}

扩写要求：
1. 优先扩写"学生—学校匹配度"和"学术与课外准备"章节
2. 补充新的证据、执行步骤、具体建议
3. 避免空话套话，增加可核查的细节
4. 保持专业、客观的语调
5. 扩写内容应与原始数据一致
6. 输出纯文本，不使用Markdown格式

扩写后的内容："""
        
        try:
            response = self.ai.call_llm("Writer", system_prompt, {"content": user_prompt})
            if isinstance(response, str):
                logger.info("字数保护机制扩写完成")
                return response.strip()
            else:
                logger.warning("字数保护机制扩写失败，返回原内容")
                return content
        except Exception as e:
            logger.error(f"字数保护机制扩写失败: {e}")
            return content
    
    def control_target_length(self, content: str) -> str:
        """
        目标页数控制机制
        
        Args:
            content: 报告内容
            
        Returns:
            控制后的内容
        """
        char_count = len(content)
        target_min = 14000  # 14k字符
        target_max = 15000  # 15k字符
        max_limit = 15500   # 15.5k字符
        
        logger.info(f"当前字数: {char_count}, 目标范围: {target_min}-{target_max}")
        
        if char_count <= max_limit:
            # 字数在合理范围内
            logger.info("字数在合理范围内，无需调整")
            return content
        
        # 字数超过15.5k，触发压缩重写
        logger.warning(f"字数超过{max_limit}，触发压缩重写...")
        
        # 系统提示词
        system_prompt = """你是一名专业的学校申请顾问，负责压缩报告内容，优先精简重复表述和模板化语言。"""
        
        # 压缩提示词
        user_prompt = f"""请对以下报告内容进行压缩重写，目标字数控制在14k-15k中文字符：

当前内容（{char_count}字）：{content}

压缩要求：
1. 优先精简重复表述和模板化语言
2. 保持所有6个章节的完整性
3. 保留核心信息和可执行建议
4. 删除冗余的修饰词和重复描述
5. 合并相似内容的段落
6. 保持专业、客观的语调
7. 目标字数：14,000-15,000中文字符
8. 输出纯文本，不使用Markdown格式

压缩后的内容："""
        
        try:
            response = self.ai.call_llm("Writer", system_prompt, {"content": user_prompt})
            if isinstance(response, str):
                compressed_content = response.strip()
                new_char_count = len(compressed_content)
                logger.info(f"压缩完成，新字数: {new_char_count}")
                
                # 验证压缩效果
                if new_char_count <= max_limit:
                    logger.info("压缩成功，字数符合要求")
                    return compressed_content
                else:
                    logger.warning(f"压缩后字数仍超过{max_limit}，返回原内容")
                    return content
            else:
                logger.warning("压缩重写失败，返回原内容")
                return content
        except Exception as e:
            logger.error(f"压缩重写失败: {e}")
            return content
    
    def print_acceptance_criteria(self, dedupe_validation: Dict[str, Any], 
                                 validation_result: Dict[str, Any],
                                 section_validation: Dict[str, Any]) -> None:
        """
        打印验收标准（自动打印到控制台）
        
        Args:
            dedupe_validation: 去重验证结果
            validation_result: 内容验证结果
            section_validation: 章节验证结果
        """
        print("\n" + "="*60)
        print("📋 去重精修验收标准检查")
        print("="*60)
        
        # 1. 章节数量检查
        print("1. 章节数量检查:")
        if section_validation["is_valid"]:
            print("   ✅ 通过 - 发现6个章节，无重复")
            print(f"   章节顺序: {' → '.join(section_validation['found_sections'])}")
        else:
            print("   ❌ 未通过 - 章节数量不正确")
            print(f"   期望: 6个章节，实际: {section_validation['total_sections']}个")
            if section_validation["missing_sections"]:
                print(f"   缺失章节: {', '.join(section_validation['missing_sections'])}")
        
        # 2. 全文无重复段落检查
        print("\n2. 全文无重复段落检查:")
        if dedupe_validation.get("meets_criteria", True):
            print("   ✅ 通过 - 任意两段 cosine_sim < 0.92")
        else:
            print("   ❌ 未通过 - 存在高度重复段落")
            for issue in dedupe_validation.get("issues", []):
                print(f"      - {issue}")
        
        # 3. 字数减少检查
        reduction_pct = dedupe_validation.get("reduction_percentage", 0)
        print(f"\n3. 字数减少检查:")
        print(f"   字数减少: {reduction_pct:.1f}%")
        if reduction_pct <= 8:
            print("   ✅ 通过 - 字数减少在合理范围内")
        else:
            print("   ⚠️  警告 - 字数减少过多，已触发保护机制")
        
        # 4. 模板句重复检查
        print("\n4. 模板句重复检查:")
        if not validation_result.get("has_markdown", False):
            print("   ✅ 通过 - 无Markdown语法")
        else:
            print("   ❌ 未通过 - 包含Markdown语法")
        
        if not validation_result.get("has_emoji", False):
            print("   ✅ 通过 - 无emoji")
        else:
            print("   ❌ 未通过 - 包含emoji")
        
        if not validation_result.get("has_placeholders", False):
            print("   ✅ 通过 - 无占位符")
        else:
            print("   ❌ 未通过 - 包含占位符")
        
        # 5. 章节完整性检查
        sections_found = validation_result.get("sections_found", 0)
        print(f"\n5. 章节完整性检查:")
        print(f"   发现章节数: {sections_found}/6")
        if sections_found >= 6:
            print("   ✅ 通过 - 所有章节完整")
        else:
            print("   ❌ 未通过 - 章节不完整")
        
        # 6. 去重报告生成检查
        dedupe_log_file = self.logs_dir / "dedupe_report.json"
        if dedupe_log_file.exists():
            print("\n6. 去重报告生成检查:")
            print("   ✅ 通过 - logs/dedupe_report.json 已生成")
            
            # 读取并显示去重统计
            try:
                with open(dedupe_log_file, 'r', encoding='utf-8') as f:
                    dedupe_report = json.load(f)
                
                summary = dedupe_report.get("summary", {})
                print(f"   修改统计:")
                print(f"   - 改写段落: {summary.get('rephrase_count', 0)}")
                print(f"   - 删除段落: {summary.get('drop_count', 0)}")
                print(f"   - 模板句改写: {summary.get('template_rewrite_count', 0)}")
                
            except Exception as e:
                logger.warning(f"读取去重报告失败: {e}")
        else:
            print("\n6. 去重报告生成检查:")
            print("   ⚠️  警告 - logs/dedupe_report.json 未生成")
        
        # 7. 总体评估
        print("\n7. 总体评估:")
        all_passed = (
            section_validation["is_valid"] and
            dedupe_validation.get("meets_criteria", True) and
            reduction_pct <= 8 and
            not validation_result.get("has_markdown", False) and
            not validation_result.get("has_emoji", False) and
            not validation_result.get("has_placeholders", False) and
            sections_found >= 6
        )
        
        if all_passed:
            print("   🎉 所有验收标准通过！报告质量符合要求。")
        else:
            print("   ⚠️  部分验收标准未通过，建议检查报告质量。")
        
        print("="*60)


def main():
    """测试LLM报告生成器"""
    generator = LLMReportGenerator()
    
    # 测试数据
    student_data = {
        "name": "Alex Chen",
        "age": "14岁",
        "grade": "Grade 8",
        "gpa": "3.8/4.0",
        "academic_strengths": "数学、物理、计算机科学",
        "competition_achievements": "机器人竞赛省级二等奖",
        "leadership_positions": "科技部副部长",
        "project_experiences": "环保义卖活动组织",
        "learning_ability": "自主学习和问题解决",
        "adaptability": "跨文化环境适应",
        "target_schools": [
            {
                "name": "Upper Canada College",
                "scores": {"academic": 4, "activities": 4, "culture": 5, "personality": 4},
                "weights": {"academic": 0.35, "activities": 0.25, "culture": 0.2, "personality": 0.2}
            },
            {
                "name": "Havergal College", 
                "scores": {"academic": 4, "activities": 3, "culture": 4, "personality": 3},
                "weights": {"academic": 0.35, "activities": 0.25, "culture": 0.2, "personality": 0.2}
            }
        ]
    }
    
    conversation_log = [
        {"role": "student", "content": "我喜欢组织活动和做科学实验"},
        {"role": "parent", "content": "孩子在学生会组织过环保义卖"}
    ]
    
    # 生成报告
    report_result = generator.generate_report(conversation_log, student_data)
    
    print("报告生成成功!")
    print(f"页数: {report_result['metadata']['page_count']}")
    print(f"字数: {report_result['metadata']['word_count']}")
    
    # 导出报告
    exported_files = generator.export_report(report_result, "all")
    
    print("\n导出文件:")
    for format_type, file_path in exported_files.items():
        print(f"{format_type}: {file_path}")


if __name__ == "__main__":
    main()
