#!/usr/bin/env python3
"""
Word报告生成器
将Markdown格式的报告转换为精美的Word文档
"""

import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("❌ 错误: 需要安装python-docx库")
    print("请运行: pip install python-docx")
    sys.exit(1)

class WordReportGenerator:
    """Word报告生成器"""
    
    def __init__(self):
        self.doc = None
        self.styles = {}
        
    def create_document(self):
        """创建Word文档"""
        self.doc = Document()
        self.setup_styles()
        return self.doc
    
    def setup_styles(self):
        """设置文档样式"""
        # 标题样式
        try:
            # 主标题样式
            title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Microsoft YaHei'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            
            # 一级标题样式
            heading1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = 'Microsoft YaHei'
            heading1_style.font.size = Pt(18)
            heading1_style.font.bold = True
            heading1_style.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色
            heading1_style.paragraph_format.space_before = Pt(12)
            heading1_style.paragraph_format.space_after = Pt(6)
            
            # 二级标题样式
            heading2_style = self.doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = 'Microsoft YaHei'
            heading2_style.font.size = Pt(14)
            heading2_style.font.bold = True
            heading2_style.font.color.rgb = RGBColor(51, 51, 51)  # 深灰色
            heading2_style.paragraph_format.space_before = Pt(8)
            heading2_style.paragraph_format.space_after = Pt(4)
            
            # 正文样式
            body_style = self.doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Microsoft YaHei'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing = 1.15
            
            # 强调文本样式
            emphasis_style = self.doc.styles.add_style('CustomEmphasis', WD_STYLE_TYPE.PARAGRAPH)
            emphasis_style.font.name = 'Microsoft YaHei'
            emphasis_style.font.size = Pt(11)
            emphasis_style.font.bold = True
            emphasis_style.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色
            
            self.styles = {
                'title': title_style,
                'heading1': heading1_style,
                'heading2': heading2_style,
                'body': body_style,
                'emphasis': emphasis_style
            }
            
        except Exception as e:
            print(f"⚠️ 样式设置警告: {e}")
            # 使用默认样式
            self.styles = {}
    
    def add_title(self, title: str):
        """添加标题"""
        if self.styles.get('title'):
            p = self.doc.add_paragraph(title, style='CustomTitle')
        else:
            p = self.doc.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(24)
            p.runs[0].font.bold = True
        
        # 添加装饰线
        self.add_decorative_line()
    
    def add_heading1(self, text: str):
        """添加一级标题"""
        if self.styles.get('heading1'):
            self.doc.add_paragraph(text, style='CustomHeading1')
        else:
            p = self.doc.add_paragraph(text)
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.bold = True
    
    def add_heading2(self, text: str):
        """添加二级标题"""
        if self.styles.get('heading2'):
            self.doc.add_paragraph(text, style='CustomHeading2')
        else:
            p = self.doc.add_paragraph(text)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True
    
    def add_paragraph(self, text: str, style: str = 'body'):
        """添加段落"""
        if self.styles.get(style):
            self.doc.add_paragraph(text, style=f'Custom{style.capitalize()}')
        else:
            p = self.doc.add_paragraph(text)
            p.runs[0].font.size = Pt(11)
    
    def add_bullet_list(self, items: List[str]):
        """添加项目符号列表"""
        for item in items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.runs[0].font.name = 'Microsoft YaHei'
            p.runs[0].font.size = Pt(11)
    
    def add_decorative_line(self):
        """添加装饰线"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("=" * 50)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(200, 200, 200)
    
    def add_section_divider(self):
        """添加章节分隔符"""
        self.doc.add_paragraph()
        self.add_decorative_line()
        self.doc.add_paragraph()
    
    def generate_word_report(self, report_content: str, student_data: Dict[str, Any], 
                           output_path: str) -> str:
        """生成Word报告"""
        try:
            # 创建文档
            self.create_document()
            
            # 解析Markdown内容
            lines = report_content.split('\n')
            
            # 添加标题
            self.add_title("🎯 私校申请策略报告")
            
            # 添加学生信息
            self.add_heading1("📋 学生概况")
            target_schools = student_data.get('target_schools', 'Upper Canada College, Havergal College, St. Andrew\'s College')
            student_info = [
                f"姓名: {student_data.get('name', 'Alex Chen')}",
                f"年龄: {student_data.get('age', '14岁')} ({student_data.get('grade', 'Grade 8')})",
                f"目标年级: Grade 9",
                f"目标学校: {target_schools}"
            ]
            self.add_bullet_list(student_info)
            
            self.add_section_divider()
            
            # 解析并添加报告内容
            self.parse_markdown_content(lines)
            
            # 添加页脚信息
            self.add_section_divider()
            footer_text = f"报告生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}\n专业顾问: 私校申请专家团队"
            self.add_paragraph(footer_text, 'emphasis')
            
            # 保存文档
            self.doc.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"❌ Word报告生成失败: {e}")
            return None
    
    def parse_markdown_content(self, lines: List[str]):
        """解析Markdown内容并添加到Word文档"""
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # 处理标题
            if line.startswith('## '):
                self.add_heading1(line[3:])
            elif line.startswith('### '):
                self.add_heading2(line[4:])
            
            # 处理列表项
            elif line.startswith('- '):
                items = []
                while i < len(lines) and lines[i].strip().startswith('- '):
                    item = lines[i].strip()[2:]
                    # 处理粗体文本
                    if '**' in item:
                        item = item.replace('**', '')
                    items.append(item)
                    i += 1
                self.add_bullet_list(items)
                continue
            
            # 处理普通段落
            elif not line.startswith('#') and not line.startswith('-') and not line.startswith('---'):
                # 处理粗体文本
                if '**' in line:
                    # 简单处理粗体文本
                    line = line.replace('**', '')
                self.add_paragraph(line)
            
            # 处理分隔符
            elif line.startswith('---'):
                self.add_section_divider()
            
            i += 1

def main():
    """测试Word报告生成器"""
    generator = WordReportGenerator()
    
    # 测试数据
    student_data = {
        "name": "Alex Chen",
        "age": "14岁",
        "grade": "Grade 8",
        "target_schools": "Upper Canada College, Havergal College, St. Andrew's College"
    }
    
    # 示例报告内容
    report_content = """# 🎯 私校申请策略报告

## 📋 学生概况
- **姓名**: Alex Chen
- **年龄**: 14岁 (Grade 8)
- **目标年级**: Grade 9
- **目标学校**: Upper Canada College, Havergal College, St. Andrew's College

---

## 🏆 核心优势分析

### 学术表现
- **GPA**: 3.8/4.0 - 优秀学术基础
- **强项领域**: 数学、物理、计算机科学
- **标准化考试**: SSAT 85th percentile
- **学术竞赛**: 机器人竞赛省级二等奖

### 领导力潜质
- **学生会职务**: 科技部副部长
- **项目经验**: 环保义卖活动组织
- **团队协作**: 跨年级合作项目
- **影响力**: 30+学生参与，800加元筹款

---

## 🎯 申请策略建议

### 1. 学术提升计划
**目标**: 全面提升学术竞争力
- ✅ **保持优势**: 继续深化STEM领域专长
- 📈 **提升空间**: 重点加强英语文学和历史
- 🎯 **考试准备**: SSAT重考目标90th percentile以上
- 📚 **学习规划**: 建立系统性的学习计划

### 2. 领导力发展路径
**目标**: 展现卓越的领导潜质
- 🚀 **深化现有**: 扩展学生会科技部影响力
- 🌟 **创新项目**: 组织更多STEM相关活动
- 🤝 **合作能力**: 建立跨年级合作项目
- 📊 **量化成果**: 记录和展示项目影响力

---

## 🎉 成功展望

### 我们的专业价值
通过我们的专业指导，Alex Chen将在申请过程中展现最佳状态：

1. **专业评估**: 全面分析学生优势和潜力
2. **策略制定**: 制定个性化的申请策略
3. **材料优化**: 优化申请材料和文书
4. **面试指导**: 提供专业的面试指导
5. **持续支持**: 全程跟踪和支持申请过程

### 预期结果
基于Alex Chen的优秀基础和我们的专业指导，**我们有信心帮助Alex Chen获得理想学校的录取**。

**让您的孩子看到希望，让我们的专业成就您的梦想！**
"""
    
    # 生成Word报告
    output_path = "test_report.docx"
    result = generator.generate_word_report(report_content, student_data, output_path)
    
    if result:
        print(f"✅ Word报告生成成功: {result}")
    else:
        print("❌ Word报告生成失败")

if __name__ == "__main__":
    main()
